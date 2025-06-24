
"""
Created on Tue May 27 21:52:07 2025

@author: joshuawolff
"""

import streamlit as st
import openai
import pandas as pd
import time
import docx
import base64


# Set API key securely
openai.api_key = st.secrets["openai_api_key"]

#Set up page and set session state to none
st.set_page_config(page_title="EOTSS Project Collaboration Advisor (Proof of Concept)", layout="centered", page_icon="📘")

if "analysis_text" not in st.session_state:
    st.session_state.analysis_text = None
    
# Hide Streamlit UI chrome
st.markdown("""
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h2 style='text-align: center;'>EOTSS Project Collaboration Advisor</h2>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align: center;'>Proof of Concept</h3", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: gray;'>Analyze inter-team dependencies in one click</p>", unsafe_allow_html=True)

with st.expander("Learn More"):
    st.write('''
        What This Tool Does
This app is designed to help you figure out which EOTSS teams you should be working with to get your project done. It reads your project plan and suggests which EOTSS teams you should collaborate with. 

In the future, it could compare project plans against documented information about what different EOTSS departments do: things like business plans, capital plans, and enterprise standards. The goal is to give you a clearer sense of who’s already working in your space, who owns what, and where there might be overlap or collaboration opportunities.

You’ll also see suggestions for projects that are similar or complementary to yours, helping you avoid duplication and spot potential partnerships.

This is an early prototype. It's not connected to the secure EOTSS environment, so it doesn’t use any internal data. That means:

Don’t upload sensitive project plans

The analysis is based only on public info from mass.gov, not internal documents like architecture review board materials

Still, it can give surprisingly helpful advice—and once it's connected to richer, internal sources, it should get even smarter. Each run takes about 4 minutes and costs around 17 cents.
    ''')

tab1, tab2 = st.tabs(["🔍 Analyze Dependencies", "📄 View Demo Plan"])


with tab1:
    st.write("Upload a project plan (e.g., summary as text) and get advice on inter-team dependencies based on embedded EOTSS group roles and standards.")
    
    # --- EOTSS GROUP DESCRIPTIONS (embedded for Option 3) ---
    group_roles = """
    Data: Enables all Enterprise Analytics components, manages governance, quality, metadata, master data, platforms, privacy, and advanced analytics.
    
    Customer Engagement: Coordinates cross-agency collaboration, chargeback processes, change management, ServiceNow tickets, and serves as the liaison with agency partners.
    
    Technology: Determines infrastructure/software needs, builds tools/platforms, supports onboarding, and manages Snowflake integrations.
    
    Digital Services: Provides UI/UX design, accessibility compliance, dashboard design standards, and constituent ID data management.
    
    MassGIS: Standardizes and analyzes geographic data, supports spatial use cases, and trains agencies on location data use.
    
    Operations: Manages accounts, access, data backup, system security, FinOps, patching, and incident response.
    
    Security & Risk: Oversees cyber risk, compliance, data classification, and post-procurement risk reviews. 
    
    Privacy: Ensures legal data access via DULAs, evaluates privacy risk, privacy impact assessments, and reviews analytics initiatives for compliance.
    
    Legal: Drafts/reviews contracts and agreements, ensures legal frameworks for data sharing and analytics.
    
    Strategy: Ensures alignment with EOTSS strategy, handles capital budgeting, change management, and prioritization.
    """
    
    accessibility_policy = """
    Component 2: Enterprise Information Technology Accessibility Policy
    
    - Applies to all IT solutions used by the Commonwealth (e.g., websites, software, digital content).
    - Requires compliance with WCAG 2.1 Level AA, Section 508, and other accessibility standards.
    - Agencies must ensure all IT systems—internally developed or procured—meet accessibility guidelines.
    - Accessibility must be integrated into planning, design, procurement, and testing.
    - If full compliance isn't possible, a mitigation plan is required.
    - Accessibility requirements must be included in procurement contracts.
    - Agencies should publish accessibility statements and contact info for support.
    """
    
    # --- Load CSV Governance Table ---
    gov_df = pd.read_csv("Expanded_IT_Governance_Table_with_Descriptions.csv")
    gov_text = ""
    for _, row in gov_df.iterrows():
         gov_text += f"Component: {row['Requirement']}\nArea: {row['Owner']}\nDescription: {row['Description']}\nRelevance: {row['Relevance']}\n\n"
    
    # --- Load CSV Governance Table ---
     
    with open("tss_deep_research.txt", "r", encoding="utf-8", errors="ignore") as file:
        gov_narrative = file.read()
    
    
    # --- USER INPUT OPTIONS ---
    input_method = st.selectbox(
        "How would you like to provide your project plan?",
        ["Use demo project plan","Upload a .doc or .docx file","Paste project plan text"]
    )
    
    
    
    project_summary = ""
    
    if input_method == "Paste project plan text":
        st.error("**Warning:** Do not enter any sensitive or protected information. "
             "This application is for demonstration purposes only and has not been assessed by the AI COE. "
             "It leverages the consumer OpenAI GPT-4 API and is not appropriate for sensitive data."
             "Please limit your input to 400 words for this small scale proof of concept")
        project_summary = st.text_area("Paste your project summary below:", height=300)
    
    elif input_method == "Upload a .doc or .docx file":
        st.error("**Warning:** Do not upload any sensitive or protected information. "
             "This application is for demonstration purposes only and has not been assessed by the AI COE. "
             "It leverages the consumer OpenAI GPT-4 API and is not appropriate for sensitive data."
             "Please limit your input to 400 words for this small scale proof of concept")
        uploaded_file = st.file_uploader("Upload your project plan (.doc or .docx)", type=["doc", "docx"])
        if uploaded_file:
            doc = docx.Document(uploaded_file)
            project_summary = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
            
            st.success("project plan loaded successfully.")
            st.text_area("Preview:", project_summary[:1000], height=150)
        
    
    elif input_method == "Use demo project plan":
        demo_path = "Example Project Plan - Climate Data Lake Project.docx"
        doc = docx.Document(demo_path)
        project_summary = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
        st.info("Using demo project plan loaded from file.")
        
    # --- Vertical Space ---
    
    st.markdown("### &nbsp;")
    
    # --- ANALYSIS ---

    if st.button("🔍 Analyze Dependencies", key="analyze_button", use_container_width=True):

        if project_summary:
            with st.spinner("Analyzing with GPT-4. This could take a few minutes..."):
                progress = st.progress(0)
                for percent in range(0, 99, 1):
                    progress.progress(percent + 1)
                    time.sleep(1)

                # Prepare full prompt
                messages = [
                    {"role": "system", "content": f"""
    You are a project planning expert for the Commonwealth of Massachusetts with particular expertise about how to navigate the bureaucracy of EOTSS.
    
    You are reviewing project plans to ensure they align with key organizational reference components:
    
    Component 1: EOTSS Group Descriptions
    {group_roles}
    
    Component 2: Enterprise IT Accessibility Policy
    {accessibility_policy}
    
    Component 3: IT Governance Reference Table
    {gov_text}
    
    Component 4: EOTSS Organization Deep Research
    {gov_narrative}
    
    """},
                    {"role": "user", "content": f"""
    Here is a summary of a project plan:
    
    {project_summary}
    
    Please evaluate this plan against the organizational components above:
    1. Identify inter-team dependencies across EOTSS groups that might be needed to contribute to the work or validate compliance with IT Governance requirements. 
       Be specific here, citing which parts of the project plan might need to involve which other EOTSS offices or comply with which Governance Requirements. 
       If the project plan lists a schedule or milestones, specificy when in the timeline each stakeholder should be looped in. 
       You do not need to stretch to anticipate dependencies for every EOTSS Office: only cite likely dependencies. 
       Don't assume that users have access to the Governance Reference Table - when citing information from the table, cite the information rather than the location in the table. Do not reference the table at all. For example, do not mention IT Governance Reference Table Component 14
       At the end of each text block describing a team the project is dependent on, include a line saying "The best way to engage [this team] is to _____". Leave the blank spot to illustrate that such contact info is important to document. 
       If the project plan cites one of these stakeholder groups as the leader for the project, then do not make suggestions about when that group should get involved with the project.
       If the project plan included a timeline with milestones or steps, then after all blocks of text describing dependencies, include a summary table listing each milestone, stage, step, or phase of the project, and the team dependencies relevant to each stage. Only do this if such a timeline was included in the original plan and list those steps/phases/milestones/stages exactly as written in the original plan. 
    2. Check for misalignment with the Accessibility Policy. Be specific about which aspects of the project plan are likely to be most subject to Accessibility policies. 
    
    Use headings to visually seperate out these two components in your output. Add a third heading called Related Projects and underneath that heading say Coming Soon
    
    """}
                ]
    
                # Call OpenAI API
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=messages,
                    temperature=0.7
                )

                #Store Session State
                st.session_state.analysis_text = response.choices[0].message.content
    
                # Once the response is received, complete the bar
                progress.progress(100)
                time.sleep(0.5)
                progress.empty()
            
            # Show results
            st.success("✅ Analysis Complete")
            #st.markdown(response.choices[0].message.content)
            
            
        else:
            st.error("No project plan text was provided.")

    if st.session_state.analysis_text is not None:
            st.markdown(st.session_state.analysis_text)
            st.download_button(
                label="Download Analysis as Text",
                data=st.session_state.analysis_text,
                file_name="project_analysis.txt",
                mime="text/plain"
        )
         
with tab2:
    st.subheader("Example (fake) Project Plan – Climate Data Lake Project")
    st.markdown("""
        # Climate Data Lake Project

## Overview and Goals

This project will create a Commonwealth Climate Data Lake in partnership with multiple agencies across the state.

The Climate Data Lake seeks to:

- Enable secure, centralized access to climate-relevant datasets from across the Commonwealth.
- Support data exploration, transformation, integration, and visualization by staff, researchers, and the public.
- Ensure appropriate public access through a mass.gov front-end while protecting privacy through data risk reviews.

This project will entail designing and building the infrastructure to host climate data, establishing classification and metadata standards, enabling governance workflows, and launching a public-facing site.

## Outcomes

The project is complete when the following deliverables have been accepted by EOTSS and participating agencies:

- The centralized Climate Data Lake is operational in the Commonwealth’s cloud environment.
- Datasets from at least three agencies have been ingested, classified, and documented.
- A public-facing mass.gov data portal is live, providing access to approved climate data.
- Privacy risk assessments are conducted for all public-facing datasets.
- Governance and access workflows are in place for internal staff and external researchers.

## Approach

The Project Manager will provide weekly status updates to key stakeholders via email.

Work will be organized into biweekly sprints, following Agile delivery practices.

Deliverables will be demoed in recurring check-ins. Feedback will be incorporated within the same sprint or used to inform the next sprint.

Close collaboration will be maintained with data owners at participating agencies to ensure proper classification, quality review, and onboarding of datasets.

## Project Milestones and Timeline

| Milestone # | Description                                           | Proposed Date |
|-------------|-------------------------------------------------------|----------------|
| 1           | Project approval – Project team receives go-ahead from EOTSS leadership and initial participating agencies | Week of 6/17   |
| 2           | Project kickoff – Kickoff meeting with stakeholders to align goals, approach, and team roles | Week of 6/24   |
| 3           | Initial agency datasets identified – Agencies deliver metadata and sample files for prioritized climate datasets | Week of 7/1    |
| 4           | Data classification and standards finalized – MDO and agencies align on data domains, tags, privacy tiers, and publishing guidelines | Week of 7/8    |
| 5           | First dataset ingested into Climate Data Lake – End-to-end ingestion of one sample dataset using agreed architecture | Week of 7/15   |
| 6           | Internal researcher access established – Internal permissions workflow built and tested for secure analyst access | Week of 7/22   |
| 7           | Public data review and risk mitigation – Privacy reviews completed on proposed public datasets; mitigation plans documented | Week of 7/29   |
| 8           | Mass.gov front end prototype complete – Mass.gov interface built to expose approved climate data to public users | Week of 8/5    |
| 9           | Stakeholder UAT – Participating agency staff and external researchers conduct user testing | Week of 8/12   |
| 10          | Launch & handoff – Launch of public portal, documentation delivery, and transition to operations | Week of 8/19   |
""")

with tab3:
    st.subheader("Learn More")
    st.markdown("""
        # What This Tool Does

This app is designed to help you figure out which EOTSS teams you should be working with to get your project done. It reads your project plan and suggests which EOTSS teams you should collaborate with. 

In the future, it could compare project plans against documented information about what different EOTSS departments do: things like business plans, capital plans, and enterprise standards. The goal is to give you a clearer sense of who’s already working in your space, who owns what, and where there might be overlap or collaboration opportunities.

You’ll also see suggestions for projects that are similar or complementary to yours, helping you avoid duplication and spot potential partnerships.

This is an early prototype. It's not connected to the secure EOTSS environment, so it doesn’t use any internal data. That means:

Don’t upload sensitive project plans

The analysis is based only on public info from mass.gov, not internal documents like architecture review board materials

Still
""")
    #try:
    #    with open("Example Project Plan - Climate Data Lake Project.pdf", "rb") as f:
    #        pdf_data = f.read()
    #
    #    base64_pdf = base64.b64encode(pdf_data).decode('utf-8')
    #    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="700px" type="application/pdf"></iframe>'
    #    st.markdown(pdf_display, unsafe_allow_html=True)

    #except FileNotFoundError:
    #    st.error("Demo PDF not found. Please ensure the file is in the correct location.")
